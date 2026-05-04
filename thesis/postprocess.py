import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(14)
H2_SIZE = Pt(12)
H3_SIZE = Pt(12)
BIB_SIZE = Pt(10)
CODE_FONT = "Courier New"
CODE_SIZE = Pt(9)
LINE_SPACING = 1.5
BLACK = RGBColor(0, 0, 0)


def force_font_xml(rPr, name):
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)
    for attr in [qn("w:asciiTheme"), qn("w:hAnsiTheme"), qn("w:eastAsiaTheme"), qn("w:cstheme")]:
        if rFonts.get(attr):
            del rFonts.attrib[attr]


def set_run_font(run, name=FONT_NAME, size=BODY_SIZE, bold=None, italic=None):
    run.font.name = name
    run.font.size = size
    run.font.color.rgb = BLACK
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    force_font_xml(run._r.get_or_add_rPr(), name)


def set_paragraph_spacing(para, before=Pt(0), after=Pt(0), line_spacing=LINE_SPACING,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, indent_left=None,
                           keep_next=False, page_break_before=False):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing
    pf.keep_with_next = keep_next
    pf.page_break_before = page_break_before
    if indent_left is not None:
        pf.left_indent = indent_left
    else:
        pf.left_indent = None
    pf.first_line_indent = None


def get_heading_level(style_name):
    m = re.search(r"(\d+)", style_name)
    return int(m.group(1)) if m else 0


def fix_paragraph(para):
    style = para.style.name

    if style.startswith("Heading"):
        level = get_heading_level(style)
        size = H1_SIZE if level == 1 else H2_SIZE
        set_paragraph_spacing(para, before=Pt(12), after=Pt(6),
                              alignment=WD_ALIGN_PARAGRAPH.LEFT,
                              keep_next=True, page_break_before=(level == 1))
        for run in para.runs:
            set_run_font(run, size=size, bold=True)

    elif style == "Source Code":
        set_paragraph_spacing(para, before=Pt(0), after=Pt(0),
                              line_spacing=1.0,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for run in para.runs:
            set_run_font(run, name=CODE_FONT, size=CODE_SIZE, bold=False)
        add_code_shading(para)

    elif style == "Bibliography":
        set_paragraph_spacing(para, before=Pt(0), after=Pt(6),
                              line_spacing=1.0,
                              alignment=WD_ALIGN_PARAGRAPH.LEFT)
        for run in para.runs:
            set_run_font(run, size=BIB_SIZE)

    elif style == "Compact":
        set_paragraph_spacing(para, before=Pt(0), after=Pt(3),
                              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                              indent_left=Cm(1.0))
        for run in para.runs:
            set_run_font(run, size=BODY_SIZE, bold=run.font.bold, italic=run.font.italic)

    else:
        set_paragraph_spacing(para, before=Pt(0), after=Pt(6))
        for run in para.runs:
            if run.font.name and "courier" in run.font.name.lower():
                set_run_font(run, name=CODE_FONT, size=CODE_SIZE, bold=run.font.bold)
            else:
                set_run_font(run, size=BODY_SIZE, bold=run.font.bold, italic=run.font.italic)


def add_code_shading(para):
    pPr = para._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")

    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for side in ["top", "bottom", "left", "right"]:
        border = pBdr.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            pBdr.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")


def fix_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "000000")
        borders.append(b)
    tblPr.append(borders)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    set_run_font(run, size=Pt(10), bold=(i == 0))


def fix_page_setup(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        sectPr = section._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:gutter"), "567")


def strip_theme_fonts(doc):
    for rPr in doc.element.iter(qn("w:rPr")):
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            for attr in [qn("w:asciiTheme"), qn("w:hAnsiTheme"),
                         qn("w:eastAsiaTheme"), qn("w:cstheme")]:
                if rFonts.get(attr) is not None:
                    del rFonts.attrib[attr]
            if not rFonts.get(qn("w:ascii")):
                rFonts.set(qn("w:ascii"), FONT_NAME)
                rFonts.set(qn("w:hAnsi"), FONT_NAME)


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs or not footer.paragraphs[0].text:
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            run = p.add_run()
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            run._r.append(fldChar1)

            run2 = p.add_run()
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = " PAGE "
            run2._r.append(instrText)

            run3 = p.add_run()
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "separate")
            run3._r.append(fldChar2)

            run4 = p.add_run("1")
            set_run_font(run4, size=Pt(10))

            run5 = p.add_run()
            fldChar3 = OxmlElement("w:fldChar")
            fldChar3.set(qn("w:fldCharType"), "end")
            run5._r.append(fldChar3)


def make_bookmark_name(text):
    clean = re.sub(r"[^\w\s]", "", text)
    clean = re.sub(r"\s+", "_", clean.strip())
    return f"_Toc_{clean[:40]}"


def add_bookmarks_to_headings(doc):
    bookmark_id = 100
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            level = get_heading_level(para.style.name)
            if level <= 2:
                bm_name = make_bookmark_name(para.text.strip())
                bm_start = OxmlElement("w:bookmarkStart")
                bm_start.set(qn("w:id"), str(bookmark_id))
                bm_start.set(qn("w:name"), bm_name)
                bm_end = OxmlElement("w:bookmarkEnd")
                bm_end.set(qn("w:id"), str(bookmark_id))
                pPr = para._p.find(qn("w:pPr"))
                if pPr is not None:
                    pPr.addnext(bm_start)
                else:
                    para._p.insert(0, bm_start)
                para._p.append(bm_end)
                bookmark_id += 1


def generate_toc(doc):
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading") and para.text.strip():
            level = get_heading_level(para.style.name)
            if level <= 2:
                headings.append((level, para.text.strip()))

    body = doc.element.body

    first_heading_idx = None
    for i, child in enumerate(body):
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None:
                pStyle = pPr.find(qn("w:pStyle"))
                if pStyle is not None and "Heading" in pStyle.get(qn("w:val"), ""):
                    first_heading_idx = i
                    break

    if first_heading_idx is None:
        return

    for sdt in body.findall(qn("w:sdt")):
        body.remove(sdt)

    toc_elements = []

    toc_title = make_paragraph("Tartalomjegyzék", size=Pt(14), bold=True,
                                alignment="center", space_before=Pt(0), space_after=Pt(12))
    toc_elements.append(toc_title)

    for level, text in headings:
        if text == "Tartalomjegyzék":
            continue
        entry = make_toc_entry(level, text)
        toc_elements.append(entry)

    page_break = OxmlElement("w:p")
    pb_run = OxmlElement("w:r")
    br_elem = OxmlElement("w:br")
    br_elem.set(qn("w:type"), "page")
    pb_run.append(br_elem)
    page_break.append(pb_run)
    toc_elements.append(page_break)

    for elem in reversed(toc_elements):
        body.insert(first_heading_idx, elem)


def make_toc_entry(level, text):
    para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    indent = OxmlElement("w:ind")
    if level == 1:
        indent.set(qn("w:left"), "0")
    elif level == 2:
        indent.set(qn("w:left"), "567")
    pPr.append(indent)

    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), "9072")
    tabs.append(tab)
    pPr.append(tabs)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "60")
    spacing.set(qn("w:after"), "60")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    para.append(pPr)

    display_text = text.upper() if level == 1 else text
    bm_name = make_bookmark_name(text)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bm_name)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "24")
    rPr.append(szCs)
    if level <= 2:
        b = OxmlElement("w:b")
        rPr.append(b)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = display_text
    run.append(t)
    hyperlink.append(run)

    tab_run = OxmlElement("w:r")
    tab_rPr = OxmlElement("w:rPr")
    tab_rFonts = OxmlElement("w:rFonts")
    tab_rFonts.set(qn("w:ascii"), FONT_NAME)
    tab_rFonts.set(qn("w:hAnsi"), FONT_NAME)
    tab_rPr.append(tab_rFonts)
    tab_sz = OxmlElement("w:sz")
    tab_sz.set(qn("w:val"), "24")
    tab_rPr.append(tab_sz)
    if level <= 2:
        tab_b = OxmlElement("w:b")
        tab_rPr.append(tab_b)
    color2 = OxmlElement("w:color")
    color2.set(qn("w:val"), "000000")
    tab_rPr.append(color2)
    u2 = OxmlElement("w:u")
    u2.set(qn("w:val"), "none")
    tab_rPr.append(u2)
    tab_run.append(tab_rPr)
    tab_char = OxmlElement("w:tab")
    tab_run.append(tab_char)
    hyperlink.append(tab_run)

    pageref_rPr_template = OxmlElement("w:rPr")
    pr_rFonts = OxmlElement("w:rFonts")
    pr_rFonts.set(qn("w:ascii"), FONT_NAME)
    pr_rFonts.set(qn("w:hAnsi"), FONT_NAME)
    pageref_rPr_template.append(pr_rFonts)
    pr_sz = OxmlElement("w:sz")
    pr_sz.set(qn("w:val"), "24")
    pageref_rPr_template.append(pr_sz)
    if level <= 2:
        pr_b = OxmlElement("w:b")
        pageref_rPr_template.append(pr_b)
    pr_color = OxmlElement("w:color")
    pr_color.set(qn("w:val"), "000000")
    pageref_rPr_template.append(pr_color)
    pr_u = OxmlElement("w:u")
    pr_u.set(qn("w:val"), "none")
    pageref_rPr_template.append(pr_u)

    def make_field_rPr():
        from copy import deepcopy
        return deepcopy(pageref_rPr_template)

    fld_begin_run = OxmlElement("w:r")
    fld_begin_run.append(make_field_rPr())
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin_run.append(fld_begin)
    hyperlink.append(fld_begin_run)

    instr_run = OxmlElement("w:r")
    instr_run.append(make_field_rPr())
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" PAGEREF {bm_name} \\h "
    instr_run.append(instr_text)
    hyperlink.append(instr_run)

    fld_sep_run = OxmlElement("w:r")
    fld_sep_run.append(make_field_rPr())
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_sep_run.append(fld_sep)
    hyperlink.append(fld_sep_run)

    fld_result_run = OxmlElement("w:r")
    fld_result_run.append(make_field_rPr())
    fld_result_t = OxmlElement("w:t")
    fld_result_t.text = "0"
    fld_result_run.append(fld_result_t)
    hyperlink.append(fld_result_run)

    fld_end_run = OxmlElement("w:r")
    fld_end_run.append(make_field_rPr())
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    fld_end_run.append(fld_end)
    hyperlink.append(fld_end_run)

    para.append(hyperlink)

    return para


def make_paragraph(text, size=BODY_SIZE, bold=False, alignment="left",
                   space_before=Pt(0), space_after=Pt(0)):
    para = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")

    jc = OxmlElement("w:jc")
    jc_val = {"left": "left", "center": "center", "right": "right", "justify": "both"}
    jc.set(qn("w:val"), jc_val.get(alignment, "left"))
    pPr.append(jc)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(int(space_before.pt * 20)))
    spacing.set(qn("w:after"), str(int(space_after.pt * 20)))
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)

    para.append(pPr)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rPr.append(rFonts)
    sz_elem = OxmlElement("w:sz")
    sz_elem.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(sz_elem)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size.pt * 2)))
    rPr.append(szCs)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)
    run.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    para.append(run)

    return para


def configure_toc_styles(doc):
    from docx.enum.style import WD_STYLE_TYPE

    styles = doc.styles

    for level in range(1, 3):
        style_name = f"toc {level}"
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = styles["Normal"]

        pf = style.paragraph_format
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        pf.line_spacing = 1.5

        if level == 1:
            style.font.bold = True
            style.font.size = Pt(12)
            style.font.all_caps = True
            pf.left_indent = Cm(0)
            pf.space_before = Pt(6)
        elif level == 2:
            style.font.bold = True
            style.font.size = Pt(12)
            pf.left_indent = Cm(1.0)
        elif level == 3:
            style.font.bold = False
            style.font.size = Pt(12)
            pf.left_indent = Cm(2.0)

        style.font.name = FONT_NAME
        style.font.color.rgb = BLACK

        rPr = style.element.get_or_add_rPr()
        force_font_xml(rPr, FONT_NAME)

        pPr = style.element.get_or_add_pPr()
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.append(tabs)
        else:
            for existing_tab in tabs.findall(qn("w:tab")):
                tabs.remove(existing_tab)

        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), "9072")
        tabs.append(tab)


def link_citations_to_bibliography(doc):
    bib_started = False
    bookmark_id_counter = 500
    ref_pattern = re.compile(r"\[(\d+)\]")

    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading") and "Irodalomjegyzék" in para.text:
            bib_started = True
            continue
        if bib_started and para.style and para.style.name.startswith("Heading"):
            break
        if bib_started and para.text.strip():
            m = ref_pattern.match(para.text.strip())
            if m:
                num = m.group(1)
                bm_name = f"_Ref_{num}"
                bm_start = OxmlElement("w:bookmarkStart")
                bm_start.set(qn("w:id"), str(bookmark_id_counter))
                bm_start.set(qn("w:name"), bm_name)
                bm_end = OxmlElement("w:bookmarkEnd")
                bm_end.set(qn("w:id"), str(bookmark_id_counter))
                pPr = para._p.find(qn("w:pPr"))
                if pPr is not None:
                    pPr.addnext(bm_start)
                else:
                    para._p.insert(0, bm_start)
                para._p.append(bm_end)
                bookmark_id_counter += 1

    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            if "Irodalomjegyzék" in para.text:
                break
            continue

        if para.style and para.style.name == "Source Code":
            continue

        full_text = para.text
        if not ref_pattern.search(full_text):
            continue

        runs_data = []
        for run in para.runs:
            runs_data.append({
                'text': run.text,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'size': run.font.size,
                'name': run.font.name,
            })

        combined = "".join(rd['text'] for rd in runs_data)
        if not ref_pattern.search(combined):
            continue

        for run in list(para.runs):
            run._r.getparent().remove(run._r)

        pos = 0
        run_idx = 0
        char_pos_in_run = 0

        def get_style_at(char_idx):
            ci = 0
            for rd in runs_data:
                if ci + len(rd['text']) > char_idx:
                    return rd
                ci += len(rd['text'])
            return runs_data[-1] if runs_data else {'bold': None, 'italic': None, 'size': BODY_SIZE, 'name': FONT_NAME}

        segments = []
        last_end = 0
        for m in ref_pattern.finditer(combined):
            if m.start() > last_end:
                segments.append(('text', combined[last_end:m.start()], last_end))
            segments.append(('ref', m.group(0), m.start()))
            last_end = m.end()
        if last_end < len(combined):
            segments.append(('text', combined[last_end:], last_end))

        for seg_type, seg_text, seg_start in segments:
            style_info = get_style_at(seg_start)
            if seg_type == 'ref':
                num = seg_text.strip("[]")
                bm_name = f"_Ref_{num}"

                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("w:anchor"), bm_name)

                run_el = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rFonts = OxmlElement("w:rFonts")
                font_name = style_info.get('name') or FONT_NAME
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)
                rPr.append(rFonts)
                sz = OxmlElement("w:sz")
                size_val = style_info.get('size') or BODY_SIZE
                sz.set(qn("w:val"), str(int(size_val.pt * 2)) if hasattr(size_val, 'pt') else "24")
                rPr.append(sz)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "000000")
                rPr.append(color)
                u = OxmlElement("w:u")
                u.set(qn("w:val"), "none")
                rPr.append(u)
                if style_info.get('bold'):
                    b = OxmlElement("w:b")
                    rPr.append(b)
                if style_info.get('italic'):
                    i = OxmlElement("w:i")
                    rPr.append(i)
                run_el.append(rPr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = seg_text
                run_el.append(t)
                hyperlink.append(run_el)
                para._p.append(hyperlink)
            else:
                run_el = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rFonts = OxmlElement("w:rFonts")
                font_name = style_info.get('name') or FONT_NAME
                rFonts.set(qn("w:ascii"), font_name)
                rFonts.set(qn("w:hAnsi"), font_name)
                rFonts.set(qn("w:cs"), font_name)
                rPr.append(rFonts)
                sz = OxmlElement("w:sz")
                size_val = style_info.get('size') or BODY_SIZE
                sz.set(qn("w:val"), str(int(size_val.pt * 2)) if hasattr(size_val, 'pt') else "24")
                rPr.append(sz)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "000000")
                rPr.append(color)
                if style_info.get('bold'):
                    b = OxmlElement("w:b")
                    rPr.append(b)
                if style_info.get('italic'):
                    i = OxmlElement("w:i")
                    rPr.append(i)
                run_el.append(rPr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = seg_text
                run_el.append(t)
                para._p.append(run_el)


def postprocess(docx_path):
    doc = Document(docx_path)

    fix_page_setup(doc)
    strip_theme_fonts(doc)

    for para in doc.paragraphs:
        fix_paragraph(para)

    for table in doc.tables:
        fix_table(table)

    configure_toc_styles(doc)
    add_bookmarks_to_headings(doc)
    generate_toc(doc)
    add_page_numbers(doc)
    add_initial_page_break(doc)

    doc.save(docx_path)
    print(f"    Post-processed: {docx_path}")


def postprocess_final(docx_path):
    doc = Document(docx_path)
    link_citations_to_bibliography(doc)
    doc.save(docx_path)
    print(f"    Citation links added: {docx_path}")


def add_initial_page_break(doc):
    body = doc.element.body
    pb_para = OxmlElement("w:p")
    pb_pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    pb_pPr.append(spacing)
    pb_para.append(pb_pPr)
    pb_run = OxmlElement("w:r")
    br_elem = OxmlElement("w:br")
    br_elem.set(qn("w:type"), "page")
    pb_run.append(br_elem)
    pb_para.append(pb_run)
    body.insert(0, pb_para)


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "--final":
        docx_path = Path(sys.argv[2])
        if not docx_path.exists():
            print(f"ERROR: {docx_path} does not exist")
            sys.exit(1)
        postprocess_final(str(docx_path))
    else:
        if len(sys.argv) < 2:
            docx_path = Path(__file__).parent / "build" / "thesis_body.docx"
        else:
            docx_path = Path(sys.argv[1])

        if not docx_path.exists():
            print(f"ERROR: {docx_path} does not exist")
            sys.exit(1)

        postprocess(str(docx_path))
